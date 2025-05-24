import numpy as np
import matplotlib.pyplot as plt
import scipy
import sounddevice as sd
import soundfile as sf

filename_intro = "SOUND_INTRO/SOUND_INTRO/sound1.wav"
data, fs = sf.read(filename_intro, dtype='float32')

print(data.dtype)
print(data)
print(fs)

# sd.play(data, fs)
# status = sd.wait()
# print("a")

# Zadanie 1

# f = open("demofile2.txt", "w")
# for i in range(len(data)):
#     f.write(str(data[i]))
# f.close()


filename01 = 'sound_L.wav'
filename02 = 'sound_R.wav'
filename03 = 'sound_mix.wav'
sf.write(filename01, data[:, 0], fs)  # Lewy kanal
sf.write(filename02, data[:, 1], fs)  # Prawy kanal
mono = np.mean(data, axis=1)  # (data[i, 0] + data[i, 1]) / 2

sf.write(filename03, mono, fs)  # średnia z obu kanałów - mono

filename01_png = 'sound_L.png'
filename02_png = 'sound_R.png'
filename03_png = 'sound_mix.png'

# Plot for left channel
plt.subplot(2, 1, 1)
plt.plot(data[:, 0])
plt.savefig(filename01_png)
# plt.show()

# Plot for right channel
plt.subplot(2, 1, 2)
plt.plot(data[:, 1])
plt.savefig(filename02_png)
# plt.show()

# Plot for mono channel
plt.figure()
plt.plot(mono)
plt.savefig(filename03_png)
# plt.show()

# Plot mono signal and its FFT
fsize = 2**8
plt.figure()
plt.subplot(2, 1, 1)
plt.plot(np.arange(0, mono.shape[0]) / fs, mono)

plt.subplot(2, 1, 2)
yf = scipy.fftpack.fft(mono, fsize)
plt.plot(np.arange(0, fs / 2, fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])))
# plt.show()

# Plot for left channel FFT
plt.figure()
plt.subplot(2, 1, 1)
plt.plot(np.arange(0, data[:, 0].shape[0]) / fs, data[:, 0])

plt.subplot(2, 1, 2)
yf = scipy.fftpack.fft(data[:, 0], fsize)
plt.plot(np.arange(0, fs, 1.0 * fs / (yf.size)), np.abs(yf))
# plt.show()

# Plot for right channel FFT
plt.figure()
plt.subplot(2, 1, 1)
plt.plot(np.arange(0, data[:, 1].shape[0]) / fs, data[:, 1])

plt.subplot(2, 1, 2)
yf = scipy.fftpack.fft(data[:, 1], fsize)
plt.plot(np.arange(0, fs, 1.0 * fs / (yf.size)), np.abs(yf))
# plt.show()

# Load sin_440Hz file for analysis
directory = "SOUND_SIN/SIN/"
filename440 = directory + "sin_440Hz.wav"
data, fs = sf.read(filename440, dtype=np.int32)

fsize = 2**8

# Plot and FFT for 440Hz sine wave
plt.figure()
plt.subplot(2, 1, 1)
plt.plot(np.arange(0, data.shape[0]) / fs, data)

plt.subplot(2, 1, 2)
yf = scipy.fftpack.fft(data, fsize)
plt.plot(np.arange(0, fs / 2, fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])))
# plt.show()

# Zadanie2

def plotAudio(Signal, Fs, TimeMargin=[0, 0.02], fsize=2**8):
    time_axis = np.arange(0, Signal.shape[0]) / Fs
    fig, axs = plt.subplots(2, 1, figsize=(10, 7))

    # Plot time domain signal
    axs[0].plot(time_axis[int(TimeMargin[0] * Fs):int(TimeMargin[1] * Fs)],
                Signal[int(TimeMargin[0] * Fs):int(TimeMargin[1] * Fs)])
    axs[0].set_xlabel("Czas [s]")
    axs[0].set_ylabel("Amplituda")
    axs[0].set_title("Sygnał w czasie")

    # Compute FFT of the signal
    yf = scipy.fftpack.fft(Signal, fsize)
    # Avoid log10(0) by adding a small epsilon to the FFT results
    epsilon = 1e-10  # Small value to prevent log10(0)
    axs[1].plot(np.arange(0, Fs / 2, Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2]) + epsilon))
    axs[1].set_xlabel("Częstotliwość [Hz]")
    axs[1].set_ylabel("Amplituda [dB]")
    axs[1].set_title("Widmo")

    plt.tight_layout()

    memfile = BytesIO()  # tworzenie bufora
    fig.savefig(memfile)  # zapis do bufora
    # plt.show()
    plt.close()
    return memfile

# save plot to docx

from docx import Document
from docx.shared import Inches
from io import BytesIO

document = Document()
document.add_heading('Lab01', 0)  # tworzenie nagłówków druga wartość to poziom nagłówka

files = ['sin60Hz.wav', 'sin440Hz.wav', 'sin8000Hz.wav']
Margins = [[0, 0.02], [0.133, 0.155]]
fsize_values = [2**8, 2**12, 2**16]

for file in files:
    document.add_heading('Plik - {}'.format(file), 2)
    for i, Margin in enumerate(Margins):
        for fsize in fsize_values:
            document.add_heading('fsize = {}'.format(fsize), 3)  # Add fsize header
            document.add_heading('Time margin {}'.format(Margin), 3)  # nagłówek sekcji, może być poziom wyżej
            fig, axs = plt.subplots(2, 1, figsize=(10, 7))  # tworzenie plota

            ############################################################
            # Execute plotAudio function and generate plots
            ############################################################
            memfile = plotAudio(data, fs, TimeMargin=Margin, fsize=fsize)

            fig.suptitle('Time margin {}'.format(Margin))  # Tytuł wykresu
            fig.tight_layout(pad=1.5)  # poprawa czytelności


            document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

            memfile.close()
            ############################################################
            # Tu dodajesz dane tekstowe - wartości, wyjście funkcji ect.
            # Optionally calculate max frequency and amplitude for each plot and add them below the graph.
            yf = scipy.fftpack.fft(data, fsize)
            max_index = np.argmax(np.abs(yf))
            max_frequency = max_index * fs / len(yf)
            document.add_paragraph(f"Max frequency: {max_frequency} Hz, Max amplitude: {np.abs(yf[max_index])}")
            ############################################################

document.save('report.docx')  # zapis do pliku
